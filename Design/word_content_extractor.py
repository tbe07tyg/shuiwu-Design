#!/usr/bin/env python3
"""
Word文档内容提取器
将Word文档转换为可被Cursor @引用的文本文件
"""

import os
import json
from docx import Document
from pathlib import Path

def extract_word_to_text(word_file_path, output_dir="word_content"):
    """将Word文档内容提取为文本文件"""
    try:
        # 创建输出目录
        os.makedirs(output_dir, exist_ok=True)
        
        # 读取Word文档
        doc = Document(word_file_path)
        
        # 获取文件名（不含扩展名）
        file_name = Path(word_file_path).stem
        
        # 创建文本文件路径
        text_file_path = os.path.join(output_dir, f"{file_name}.md")
        
        # 提取内容并写入Markdown文件
        with open(text_file_path, 'w', encoding='utf-8') as f:
            f.write(f"# Word文档内容: {file_name}\n\n")
            f.write(f"**原文档路径**: {word_file_path}\n")
            f.write(f"**生成时间**: {Path(word_file_path).stat().st_mtime}\n\n")
            
            # 基本信息
            f.write("## 📊 文档信息\n\n")
            f.write(f"- 段落数量: {len(doc.paragraphs)}\n")
            f.write(f"- 表格数量: {len(doc.tables)}\n")
            f.write(f"- 总字符数: {sum(len(p.text) for p in doc.paragraphs)}\n\n")
            
            # 提取文档内容
            f.write("## 📖 文档内容\n\n")
            
            current_heading_level = 0
            for para in doc.paragraphs:
                if para.text.strip():
                    style_name = para.style.name
                    text = para.text.strip()
                    
                    # 处理标题
                    if style_name == 'Title':
                        f.write(f"# {text}\n\n")
                    elif 'Heading 1' in style_name:
                        f.write(f"## {text}\n\n")
                    elif 'Heading 2' in style_name:
                        f.write(f"### {text}\n\n")
                    elif 'Heading 3' in style_name:
                        f.write(f"#### {text}\n\n")
                    elif 'List Bullet' in style_name:
                        f.write(f"- {text}\n")
                    elif 'List Number' in style_name:
                        f.write(f"1. {text}\n")
                    else:
                        f.write(f"{text}\n\n")
            
            # 提取表格数据
            if doc.tables:
                f.write("## 📋 表格数据\n\n")
                for i, table in enumerate(doc.tables, 1):
                    f.write(f"### 表格 {i}\n\n")
                    
                    # 提取表格内容为Markdown表格
                    if table.rows:
                        # 表头
                        header_row = table.rows[0]
                        headers = [cell.text.strip() for cell in header_row.cells]
                        f.write("| " + " | ".join(headers) + " |\n")
                        f.write("| " + " | ".join(["---"] * len(headers)) + " |\n")
                        
                        # 数据行
                        for row in table.rows[1:]:
                            row_data = [cell.text.strip() for cell in row.cells]
                            f.write("| " + " | ".join(row_data) + " |\n")
                        f.write("\n")
        
        return text_file_path
        
    except Exception as e:
        print(f"提取失败: {str(e)}")
        return None

def batch_extract_words():
    """批量提取工作区中的所有Word文档"""
    word_files = []
    
    # 查找所有Word文档
    for root, dirs, files in os.walk("."):
        for file in files:
            if file.endswith('.docx') and not file.startswith('~'):
                full_path = os.path.join(root, file)
                word_files.append(full_path)
    
    print(f"🔍 找到 {len(word_files)} 个Word文档")
    
    extracted_files = []
    for word_file in word_files:
        print(f"📄 正在处理: {word_file}")
        text_file = extract_word_to_text(word_file)
        if text_file:
            extracted_files.append(text_file)
            print(f"✅ 已转换: {text_file}")
    
    print(f"\n🎉 转换完成！生成了 {len(extracted_files)} 个可@引用的文件:")
    for file in extracted_files:
        print(f"   📝 {file}")
    
    print("\n💡 现在您可以在Cursor中使用 @word_content/文件名.md 来引用Word文档内容了！")
    
    return extracted_files

if __name__ == "__main__":
    batch_extract_words() 