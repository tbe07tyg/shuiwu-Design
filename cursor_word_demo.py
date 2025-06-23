#!/usr/bin/env python3
"""
Cursor环境中的Word文档识别功能演示
直接在当前对话中使用，无需额外配置
"""

import os
from docx import Document
import json

def list_word_documents():
    """列出当前目录中的所有Word文档"""
    word_files = []
    for file in os.listdir('.'):
        if file.endswith('.docx') and not file.startswith('~'):
            word_files.append(file)
    return word_files

def analyze_word_document(filename):
    """分析Word文档内容"""
    try:
        doc = Document(filename)
        
        analysis = {
            "文档名称": filename,
            "基本信息": {
                "段落数": len(doc.paragraphs),
                "表格数": len(doc.tables),
                "总字符数": sum(len(p.text) for p in doc.paragraphs)
            },
            "内容预览": [],
            "表格数据": []
        }
        
        # 提取前5个非空段落
        content_count = 0
        for para in doc.paragraphs:
            if para.text.strip() and content_count < 5:
                analysis["内容预览"].append({
                    "样式": para.style.name,
                    "内容": para.text.strip()
                })
                content_count += 1
        
        # 提取表格数据
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            
            analysis["表格数据"].append({
                "表格序号": i + 1,
                "尺寸": f"{len(table.rows)}行 x {len(table.columns)}列",
                "数据": table_data[:3]  # 只显示前3行
            })
        
        return analysis
    
    except Exception as e:
        return {"错误": f"无法分析文档: {str(e)}"}

def demonstrate_word_processing():
    """演示Word文档处理功能"""
    print("🔍 Cursor中的Word文档识别功能演示")
    print("=" * 60)
    
    # 1. 列出所有Word文档
    word_files = list_word_documents()
    
    if not word_files:
        print("❌ 当前目录中没有找到Word文档")
        print("💡 您可以:")
        print("   1. 上传一个Word文档到当前目录")
        print("   2. 或者请求我创建一个演示文档")
        return
    
    print(f"📄 发现 {len(word_files)} 个Word文档:")
    for i, file in enumerate(word_files, 1):
        size = os.path.getsize(file) / 1024  # KB
        print(f"   {i}. {file} ({size:.1f} KB)")
    
    print("\n" + "="*60)
    
    # 2. 分析第一个文档
    target_file = word_files[0]
    print(f"📖 正在分析: {target_file}")
    print("-" * 40)
    
    analysis = analyze_word_document(target_file)
    
    # 3. 显示分析结果
    if "错误" in analysis:
        print(f"❌ {analysis['错误']}")
        return
    
    print("✅ 分析完成！")
    print(json.dumps(analysis, ensure_ascii=False, indent=2))
    
    print("\n" + "="*60)
    print("🎯 在Cursor中使用Word文档识别功能的方法:")
    print("   1. 直接在对话中请求: '分析这个Word文档'")
    print("   2. 具体操作请求: '提取文档中的表格数据'") 
    print("   3. 编辑请求: '在文档中添加一个新段落'")
    print("   4. 格式化请求: '将标题改为粗体'")

if __name__ == "__main__":
    demonstrate_word_processing() 