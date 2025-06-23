#!/usr/bin/env python3
"""
Cursor Word文档助手
直接在Cursor对话中提供Word文档识别和处理功能
无需复杂的MCP配置，直接可用
"""

import os
import json
from docx import Document
from pathlib import Path

class CursorWordHelper:
    """Cursor环境下的Word文档处理助手"""
    
    def __init__(self, workspace_path="."):
        self.workspace_path = workspace_path
        
    def find_word_documents(self):
        """在工作区中查找所有Word文档"""
        word_files = []
        for root, dirs, files in os.walk(self.workspace_path):
            for file in files:
                if file.endswith('.docx') and not file.startswith('~'):
                    full_path = os.path.join(root, file)
                    rel_path = os.path.relpath(full_path, self.workspace_path)
                    size_kb = os.path.getsize(full_path) / 1024
                    word_files.append({
                        'name': file,
                        'path': rel_path,
                        'full_path': full_path,
                        'size_kb': round(size_kb, 1)
                    })
        return word_files
    
    def analyze_document(self, file_path):
        """深度分析Word文档"""
        try:
            doc = Document(file_path)
            
            analysis = {
                "文档路径": file_path,
                "基本信息": {
                    "段落总数": len(doc.paragraphs),
                    "表格总数": len(doc.tables),
                    "图片总数": self._count_images(doc),
                    "总字符数": sum(len(p.text) for p in doc.paragraphs),
                    "非空段落数": len([p for p in doc.paragraphs if p.text.strip()])
                },
                "文档结构": self._extract_structure(doc),
                "内容摘要": self._extract_summary(doc),
                "表格数据": self._extract_tables(doc),
                "样式统计": self._analyze_styles(doc)
            }
            
            return analysis
            
        except Exception as e:
            return {"错误": f"无法分析文档 {file_path}: {str(e)}"}
    
    def _count_images(self, doc):
        """统计文档中的图片数量"""
        count = 0
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if hasattr(run, '_element') and run._element.xpath('.//a:blip'):
                    count += 1
        return count
    
    def _extract_structure(self, doc):
        """提取文档结构"""
        structure = []
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                structure.append({
                    "段落号": i + 1,
                    "样式": para.style.name,
                    "级别": self._get_heading_level(para.style.name),
                    "内容": para.text.strip()[:100] + ("..." if len(para.text.strip()) > 100 else "")
                })
        return structure
    
    def _get_heading_level(self, style_name):
        """获取标题级别"""
        if 'Heading' in style_name:
            try:
                return int(style_name.split()[-1])
            except:
                return 0
        elif style_name == 'Title':
            return 0
        return None
    
    def _extract_summary(self, doc):
        """提取文档摘要"""
        summary = {
            "标题": "",
            "主要内容": [],
            "关键信息": []
        }
        
        # 提取标题
        for para in doc.paragraphs:
            if para.style.name in ['Title', 'Heading 1']:
                summary["标题"] = para.text.strip()
                break
        
        # 提取主要段落
        content_paras = [p for p in doc.paragraphs if p.text.strip() and 
                        p.style.name not in ['Title', 'Heading 1', 'Heading 2', 'Heading 3']]
        summary["主要内容"] = [p.text.strip() for p in content_paras[:3]]
        
        # 提取关键信息（数字、日期等）
        import re
        all_text = " ".join([p.text for p in doc.paragraphs])
        numbers = re.findall(r'\d+', all_text)
        summary["关键信息"] = {
            "数字统计": len(numbers),
            "包含数字": numbers[:10] if numbers else []
        }
        
        return summary
    
    def _extract_tables(self, doc):
        """提取表格数据"""
        tables_data = []
        for i, table in enumerate(doc.tables):
            table_info = {
                "表格序号": i + 1,
                "尺寸": f"{len(table.rows)}行 x {len(table.columns)}列",
                "表头": [],
                "数据预览": []
            }
            
            if table.rows:
                # 提取表头
                header_row = table.rows[0]
                table_info["表头"] = [cell.text.strip() for cell in header_row.cells]
                
                # 提取数据预览（前5行）
                for row_idx, row in enumerate(table.rows[1:6], 1):
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_info["数据预览"].append({
                        f"第{row_idx}行": row_data
                    })
            
            tables_data.append(table_info)
        
        return tables_data
    
    def _analyze_styles(self, doc):
        """分析文档样式使用情况"""
        style_count = {}
        for para in doc.paragraphs:
            style_name = para.style.name
            style_count[style_name] = style_count.get(style_name, 0) + 1
        
        return dict(sorted(style_count.items(), key=lambda x: x[1], reverse=True))
    
    def search_in_document(self, file_path, search_term):
        """在文档中搜索特定内容"""
        try:
            doc = Document(file_path)
            results = []
            
            for i, para in enumerate(doc.paragraphs):
                if search_term.lower() in para.text.lower():
                    results.append({
                        "段落号": i + 1,
                        "内容": para.text.strip(),
                        "样式": para.style.name
                    })
            
            return {
                "搜索词": search_term,
                "找到结果": len(results),
                "匹配内容": results
            }
            
        except Exception as e:
            return {"错误": f"搜索失败: {str(e)}"}
    
    def extract_text_only(self, file_path):
        """仅提取文档的纯文本内容"""
        try:
            doc = Document(file_path)
            text_content = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text.strip())
            
            return {
                "纯文本内容": text_content,
                "总段落数": len(text_content),
                "总字符数": sum(len(text) for text in text_content)
            }
            
        except Exception as e:
            return {"错误": f"提取文本失败: {str(e)}"}

def main():
    """演示功能"""
    print("🔍 Cursor Word文档助手")
    print("=" * 50)
    
    helper = CursorWordHelper()
    
    # 1. 查找Word文档
    word_docs = helper.find_word_documents()
    
    if not word_docs:
        print("❌ 未找到Word文档")
        return
    
    print(f"📄 找到 {len(word_docs)} 个Word文档:")
    for doc in word_docs:
        print(f"   📝 {doc['name']} ({doc['size_kb']} KB)")
        print(f"      路径: {doc['path']}")
    
    # 2. 分析第一个文档
    if word_docs:
        target_doc = word_docs[0]
        print(f"\n🔍 正在分析: {target_doc['name']}")
        print("-" * 40)
        
        analysis = helper.analyze_document(target_doc['full_path'])
        print("✅ 分析完成！")
        print(json.dumps(analysis, ensure_ascii=False, indent=2))

if __name__ == "__main__":
    main() 